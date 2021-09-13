# -*- coding: utf-8 -*-

import base64
import docx
import logging
import filetype

from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from .quickstart import download_files
from tempfile import gettempdir

from odoo import api, fields, models

_logger = logging.getLogger(__name__)


class TransferenciaWordReportOut(models.Model):
    _name = 'transferencia.word.report.out'

    binary_word_report = fields.Binary('Reporte en Word', readonly=True)
    binary_word_report_filename = fields.Char('Reporte en Word')


class TransferenciaWordReport(models.Model):
    _name = 'transferencia.word.report'
    _rec_name = 'code'
    _order = 'date_start,code'

    @api.model
    def default_get(self, fields_list):
        res = super(TransferenciaWordReport, self).default_get(fields_list)
        values = {'code': self.env['ir.sequence'].next_by_code('transferencia.sequence')}
        res.update(values)
        return res

    code = fields.Char('Código')
    date_start = fields.Datetime('Inicio')
    date_end = fields.Datetime('Fin')
    skip = fields.Boolean('¿Saltar Evaluación de Token?')
    lot_number = fields.Char('Número de lote')

    @staticmethod
    def reduce_size(tmp, rec):
        foo = Image.open('{}/{}'.format(tmp, rec))
        foo = foo.resize((590, 790), Image.ANTIALIAS)
        foo.save('{}/{}_{}'.format(tmp, 'RESIZE', rec), optimize=True, quality=95)

    @staticmethod
    def get_file_type(file_location):
        kind = filetype.guess(file_location)
        if kind is None:
            print('Cannot guess file type!')
            return
        print('File extension: %s' % kind.extension)
        return kind.extension, kind.mime

    def _convert_http(self, doc, field, title, tmp):
        self.ensure_one()
        doc.add_paragraph(title, 'heading 1')
        fields_index = field.split(',')
        file_ids = []
        file_names = []
        for f in fields_index:
            ubi = f.find('=')
            file_id = f[ubi+1:]
            file_ids.append(file_id)
            file_names.append(file_id + '.jpg')

        download_files(file_ids, file_names, self.skip)
        for rec in file_names:
            try:
                doc.add_picture('{}/{}'.format(tmp, rec), width=docx.shared.Cm(15), height=docx.shared.Cm(12))
            except Exception as e:
                _logger.error(u'Error al agregar Imagen: {}, {}'.format(rec, e))
                location_file = '{}/{}'.format(tmp, rec)
                extension, mime = self.get_file_type(location_file)
                if extension == 'jpg':
                    self.reduce_size(tmp, rec)
                    try:
                        doc.add_picture('{}/{}_{}'.format(tmp, 'RESIZE', rec), width=docx.shared.Cm(15), height=docx.shared.Cm(12))
                    except Exception as e:
                        _logger.error(u'Excepción al intentar agregar Imagen comprimida: {}, {}'.format(rec, e))
                if extension == 'pdf':
                    images = convert_from_path(location_file, 500, poppler_path='C:/Users/MILTON/Desktop/poppler-0.68.0/bin')
                    for i in range(len(images)):
                        # Save pages as images in the pdf
                        images[i].save('{}/{}'.format(tmp, rec) + str(i) + '.jpg', 'JPEG')
                        try:
                            doc.add_picture('{}/{}'.format(tmp, rec) + str(i) + '.jpg', width=docx.shared.Cm(15),
                                            height=docx.shared.Cm(20))
                        except Exception as e:
                            _logger.error(u'NUEVAMENTE Excepción al agregar Imagen: {}'.format(e))

    def action_report_word(self):
        self.ensure_one()
        tmp = gettempdir()
        Transferencia = self.env['transferencia.form']
        date_start = fields.Datetime.from_string(self.date_start)
        date_end = fields.Datetime.from_string(self.date_end)
        domain = [
            ('marca_temporal', '>=', date_start),
            ('marca_temporal', '<=', date_end),
            ('lot_number', '=', self.lot_number)
        ]
        transferencia_ids = Transferencia.search(domain)
        for trans in transferencia_ids:
            doc = Document()
            doc.add_paragraph(u'ID Código: {}'.format(trans.codigo_mercado), 'heading 1')
            doc.add_paragraph(u'Región: {}'.format(trans.departamento), 'heading 1')
            doc.add_paragraph(u'Provincia: {}'.format(trans.provincia), 'heading 1')
            doc.add_paragraph(u'Distrito: {}'.format(trans.distrito), 'heading 1')
            doc.add_paragraph(u'Nombre: {}'.format(trans.nombre_mercado), 'heading 1')
            doc.add_paragraph('')

            self._convert_http(doc, trans.f0, 'Foto de entrega del acta', tmp)

            self._convert_http(doc, trans.f11, 'Foto de cumplimiento medida 1', tmp)
            self._convert_http(doc, trans.f12, 'Foto de NO cumplimiento medida 1', tmp)
            self._convert_http(doc, trans.f21, 'Foto de cumplimiento medida 2', tmp)
            self._convert_http(doc, trans.f22, 'Foto de NO cumplimiento medida 2', tmp)
            self._convert_http(doc, trans.f31, 'Foto de cumplimiento medida 3', tmp)
            self._convert_http(doc, trans.f32, 'Foto de NO cumplimiento medida 3', tmp)
            self._convert_http(doc, trans.f41, 'Foto de cumplimiento medida 4', tmp)
            self._convert_http(doc, trans.f42, 'Foto de NO cumplimiento medida 4', tmp)

            self._convert_http(doc, trans.f51, 'Foto de cumplimiento medida 5', tmp)
            self._convert_http(doc, trans.f52, 'Foto de NO cumplimiento medida 5', tmp)
            self._convert_http(doc, trans.f61, 'Foto de cumplimiento medida 6', tmp)
            self._convert_http(doc, trans.f62, 'Foto de NO cumplimiento medida 6', tmp)
            self._convert_http(doc, trans.f71, 'Foto de cumplimiento medida 7', tmp)
            self._convert_http(doc, trans.f72, 'Foto de NO cumplimiento medida 7', tmp)
            self._convert_http(doc, trans.f81, 'Foto de cumplimiento medida 8', tmp)
            self._convert_http(doc, trans.f82, 'Foto de NO cumplimiento medida 8', tmp)

            self._convert_http(doc, trans.f91, 'Foto de cumplimiento medida 9', tmp)
            self._convert_http(doc, trans.f92, 'Foto de NO cumplimiento medida 9', tmp)
            self._convert_http(doc, trans.f101, 'Foto de cumplimiento medida 10', tmp)
            self._convert_http(doc, trans.f102, 'Foto de NO cumplimiento medida 10', tmp)
            self._convert_http(doc, trans.f111, 'Foto de cumplimiento medida 11', tmp)
            self._convert_http(doc, trans.f112, 'Foto de NO cumplimiento medida 11', tmp)
            self._convert_http(doc, trans.f121, 'Foto de cumplimiento medida 12', tmp)
            self._convert_http(doc, trans.f122, 'Foto de NO cumplimiento medida 12', tmp)

            self._convert_http(doc, trans.f131, 'Foto de cumplimiento medida 13', tmp)
            self._convert_http(doc, trans.f132, 'Foto de NO cumplimiento medida 13', tmp)
            self._convert_http(doc, trans.f141, 'Foto de cumplimiento medida 14', tmp)
            self._convert_http(doc, trans.f142, 'Foto de NO cumplimiento medida 14', tmp)
            self._convert_http(doc, trans.f151, 'Foto de cumplimiento medida 15', tmp)
            self._convert_http(doc, trans.f152, 'Foto de NO cumplimiento medida 15', tmp)
            self._convert_http(doc, trans.f161, 'Foto de cumplimiento medida 16', tmp)
            self._convert_http(doc, trans.f162, 'Foto de NO cumplimiento medida 16', tmp)

            self._convert_http(doc, trans.f171, 'Foto de cumplimiento medida 17', tmp)
            self._convert_http(doc, trans.f172, 'Foto de NO cumplimiento medida 17', tmp)
            self._convert_http(doc, trans.f181, 'Foto de cumplimiento medida 18', tmp)
            self._convert_http(doc, trans.f182, 'Foto de NO cumplimiento medida 18', tmp)
            self._convert_http(doc, trans.f191, 'Foto de cumplimiento medida 19', tmp)
            self._convert_http(doc, trans.f192, 'Foto de NO cumplimiento medida 19', tmp)
            self._convert_http(doc, trans.f201, 'Foto de cumplimiento medida 20', tmp)
            self._convert_http(doc, trans.f202, 'Foto de NO cumplimiento medida 20', tmp)

            word_filename = '{}/{}.docx'.format(tmp, trans.codigo_mercado)
            doc.save(word_filename)
            word_filename2 = '{}/{}.docx'.format('E:/MINSA/OS9/META6/Consolidado_Mercados_9_2_ORP', trans.codigo_mercado)
            doc.save(word_filename2)

        _logger.info(u'SE COMPLETÓ LA DESCARGA DE {} ARCHIVOS WORD'.format(len(transferencia_ids)))

        valores_luneros = {'binary_word_report_filename': u'SE COMPLETÓ LA DESCARGA DE {} ARCHIVOS WORD'.format(len(transferencia_ids))}

        act_id = self.env['transferencia.word.report.out'].create(valores_luneros)
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'transferencia.word.report.out',
            'res_id': act_id.id,
            'view_type': 'form',
            'view_mode': 'form',
            'context': self.env.context,
            'target': 'new',
        }
